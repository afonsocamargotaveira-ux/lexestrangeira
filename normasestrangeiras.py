import streamlit as st
import requests
from bs4 import BeautifulSoup
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import urllib.parse
import re

# ─── Page config ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="LexGlobal – Legislação Estrangeira",
    page_icon="⚖️",
    layout="wide",
)

# ─── Countries & URL templates ──────────────────────────────────────────────
COUNTRIES = {
    "Austrália": {
        "url": 'https://www.legislation.gov.au/search/text(%22{term}%22,nameAndText,contains)/pointintime(Latest)',
        "term_param": "text",
        "flag": "🇦🇺",
    },
    "Nova Zelândia": {
        "url": "https://www.legislation.govt.nz/items/?search_term={term}&search_field=title&per_page=20&search_for=in_force&legislation_status=in_force",
        "term_param": "search_term",
        "flag": "🇳🇿",
    },
    "Reino Unido": {
        "url": "https://www.legislation.gov.uk/all?text={term}",
        "term_param": "text",
        "flag": "🇬🇧",
    },
    "Canadá": {
        "url": "https://www.parl.ca/LegisInfo/en/bills?keywords={term}",
        "term_param": "keywords",
        "flag": "🇨🇦",
    },
    "Estados Unidos": {
        "url": 'https://www.govinfo.gov/app/search/%7B%22historical%22%3Atrue%2C%22offset%22%3A0%2C%22query%22%3A%22collection%3A(STATUTE%20OR%20USCODE)%20AND%20publishdate%3Arange(%2C2026-06-03)%20%20AND%20%20content%3A({term})%22%7D',
        "term_param": "content",
        "flag": "🇺🇸",
        "raw": True,
    },
    "Irlanda": {
        "url": "https://www.irishstatutebook.ie/eli/ResultsTitle.html?q={term}",
        "term_param": "q",
        "flag": "🇮🇪",
    },
    "Índia": {
        "url": "https://www.indiacode.nic.in/handle/123456789/1362/simple-search?query={term}&btngo=&searchradio=acts",
        "term_param": "query",
        "flag": "🇮🇳",
    },
}


def build_url(country_key: str, term: str) -> str:
    info = COUNTRIES[country_key]
    encoded = urllib.parse.quote_plus(term)
    if info.get("raw"):
        return info["url"].replace("{term}", term)
    if country_key == "Austrália":
        encoded = urllib.parse.quote(term, safe="")  # espaços viram %20
    return info["url"].replace("{term}", encoded)


# ─── CSS ────────────────────────────────────────────────────────────────────
st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&family=DM+Sans:wght@300;400;500;600&display=swap');

/* ── Global reset ── */
html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
}

/* ── Hide Streamlit chrome ── */
#MainMenu, footer, header {visibility: hidden;}
.stDeployButton {display: none;}

/* ── App background ── */
.stApp {
    background: #FFFFFF;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #FF6B00 0%, #FF8C00 100%) !important;
    border-right: none;
}
[data-testid="stSidebar"] * {color: white !important;}
[data-testid="stSidebar"] .stMultiSelect [data-baseweb="tag"] {
    background-color: rgba(255,255,255,0.25) !important;
}
[data-testid="stSidebar"] label {font-weight: 600 !important; letter-spacing: 0.05em;}
[data-testid="stSidebar"] .stSelectbox > div > div,
[data-testid="stSidebar"] .stMultiSelect > div > div {
    background: rgba(255,255,255,0.18) !important;
    border: 1.5px solid rgba(255,255,255,0.4) !important;
    color: white !important;
    border-radius: 8px !important;
}
[data-testid="stSidebar"] .stTextInput > div > div > input {
    background: white !important;
    border: 1.5px solid rgba(255,255,255,0.6) !important;
    color: #1A1A1A !important;
    border-radius: 8px !important;
}
[data-testid="stSidebar"] .stTextInput input::placeholder {color: #AAAAAA !important;}
[data-testid="stSidebar"] .stDateInput input {
    background: white !important;
    border: 1.5px solid rgba(255,255,255,0.6) !important;
    color: #1A1A1A !important;
    border-radius: 8px !important;
}
[data-testid="stSidebar"] button {
    background: white !important;
    color: #FF6B00 !important;
    font-weight: 700 !important;
    border-radius: 8px !important;
    border: none !important;
    letter-spacing: 0.05em;
}
[data-testid="stSidebar"] button:hover {
    background: #FFF3E0 !important;
    color: #E65000 !important;
}
[data-testid="stSidebar"] .stCheckbox span {color: white !important;}

/* ── Main header ── */
.lex-header {
    background: linear-gradient(135deg, #FF6B00 0%, #FF8C00 60%, #FFA040 100%);
    border-radius: 16px;
    padding: 32px 40px;
    margin-bottom: 28px;
    display: flex;
    align-items: center;
    gap: 24px;
    box-shadow: 0 8px 32px rgba(255,107,0,0.22);
}
.lex-logo-wrap {
    display: flex;
    align-items: center;
    gap: 14px;
}
.lex-icon {
    width: 58px; height: 58px;
    background: white;
    border-radius: 14px;
    display: flex; align-items: center; justify-content: center;
    font-size: 30px;
    box-shadow: 0 4px 16px rgba(0,0,0,0.12);
    flex-shrink: 0;
}
.lex-title {
    font-family: 'Playfair Display', serif;
    font-size: 2.6rem;
    font-weight: 900;
    color: white;
    letter-spacing: -0.02em;
    line-height: 1;
    margin: 0;
}
.lex-title span {color: rgba(255,255,255,0.6); font-weight: 700; font-size: 1.5rem;}
.lex-subtitle {
    font-size: 0.95rem;
    color: rgba(255,255,255,0.85);
    margin-top: 4px;
    font-weight: 300;
    letter-spacing: 0.04em;
    text-transform: uppercase;
}
.lex-divider {
    width: 2px; height: 60px;
    background: rgba(255,255,255,0.35);
    border-radius: 2px;
    margin: 0 8px;
}
.lex-tagline {
    color: rgba(255,255,255,0.92);
    font-size: 1.05rem;
    font-weight: 300;
    line-height: 1.5;
    max-width: 420px;
}

/* ── Country cards ── */
.country-card {
    background: #FFFFFF;
    border: 2px solid #F0F0F0;
    border-radius: 14px;
    padding: 20px 22px;
    margin-bottom: 14px;
    transition: all 0.2s ease;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}
.country-card:hover {
    border-color: #FF6B00;
    box-shadow: 0 4px 20px rgba(255,107,0,0.12);
}
.country-name {
    font-size: 1.1rem;
    font-weight: 700;
    color: #1A1A1A;
    margin-bottom: 6px;
    display: flex; align-items: center; gap: 8px;
}
.country-url {
    font-size: 0.82rem;
    color: #FF6B00;
    word-break: break-all;
    font-family: 'DM Mono', monospace;
    text-decoration: none;
}
.country-url:hover {color: #E65000;}

/* ── Section titles ── */
.section-title {
    font-family: 'Playfair Display', serif;
    font-size: 1.35rem;
    font-weight: 700;
    color: #1A1A1A;
    margin: 0 0 18px 0;
    padding-bottom: 10px;
    border-bottom: 3px solid #FF6B00;
    display: inline-block;
}

/* ── Info chips ── */
.chip {
    display: inline-flex; align-items: center; gap: 6px;
    background: #FFF3E0;
    color: #FF6B00;
    border-radius: 20px;
    padding: 4px 12px;
    font-size: 0.8rem;
    font-weight: 600;
    margin: 0 4px 4px 0;
    border: 1.5px solid #FFD9B0;
}

/* ── Download button ── */
.stDownloadButton > button {
    background: linear-gradient(135deg, #FF6B00, #FF8C00) !important;
    color: white !important;
    font-weight: 700 !important;
    border-radius: 10px !important;
    border: none !important;
    padding: 12px 28px !important;
    font-size: 1rem !important;
    letter-spacing: 0.04em;
    box-shadow: 0 4px 16px rgba(255,107,0,0.3) !important;
    transition: all 0.2s !important;
}
.stDownloadButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 24px rgba(255,107,0,0.4) !important;
}

/* ── Alerts ── */
.stAlert {border-radius: 10px !important;}

/* ── Spinner ── */
.stSpinner > div {border-top-color: #FF6B00 !important;}

/* ── Metric ── */
[data-testid="metric-container"] {
    background: #FFF8F3;
    border: 2px solid #FFE0C0;
    border-radius: 12px;
    padding: 12px 16px;
}
</style>
""",
    unsafe_allow_html=True,
)

# ─── Header / Logo ───────────────────────────────────────────────────────────
st.markdown(
    """
<div class="lex-header">
  <div class="lex-logo-wrap">
    <div class="lex-icon">⚖️</div>
    <div>
      <div class="lex-title">Lex<span>Global</span></div>
      <div class="lex-subtitle">Pesquisa de Legislação Internacional</div>
    </div>
  </div>
  <div class="lex-divider"></div>
  <div class="lex-tagline">
    Acesse bases de dados legislativas de parlamentos de língua inglesa em um único lugar.
    Pesquise, compare e exporte resultados com facilidade.
  </div>
</div>
""",
    unsafe_allow_html=True,
)

# ─── Sidebar ─────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 🔍 Parâmetros de Busca")
    st.markdown("---")

    search_term = st.text_input(
        "Termo de busca",
        placeholder="Ex.: police, education, tax...",
    )

    st.markdown("#### 🌍 Países")
    all_countries = list(COUNTRIES.keys())
    select_all = st.checkbox("Selecionar todos", value=True)
    if select_all:
        selected_countries = all_countries
    else:
        selected_countries = st.multiselect(
            "Escolha os países",
            options=all_countries,
            format_func=lambda c: f"{COUNTRIES[c]['flag']} {c}",
        )

    st.markdown("---")
    search_btn = st.button("🔎 Realizar Busca", use_container_width=True)
    st.markdown("---")
    st.markdown(
        "<small style='color:rgba(255,255,255,0.55)'>LexGlobal © 2025 · v1.0</small>",
        unsafe_allow_html=True,
    )

# ─── Main content ─────────────────────────────────────────────────────────────
if not search_btn:
    st.markdown('<p class="section-title">Como usar</p>', unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        st.info("**1. Digite o termo** de busca no painel lateral (ex.: *police*, *education*).")
    with col2:
        st.info("**2. Selecione os países** cujas legislações deseja pesquisar.")
    with col3:
        st.info("**3. Clique em Realizar Busca** e exporte os resultados em Word.")

    st.markdown('<br>', unsafe_allow_html=True)
    st.markdown('<p class="section-title">Bases de dados disponíveis</p>', unsafe_allow_html=True)

    cols = st.columns(2)
    for i, (country, info) in enumerate(COUNTRIES.items()):
        with cols[i % 2]:
            sample_url = build_url(country, "school")
            st.markdown(
                f"""
                <div class="country-card">
                  <div class="country-name">{info['flag']} {country}</div>
                  <a href="{sample_url}" target="_blank" class="country-url">{sample_url[:70]}{"…" if len(sample_url) > 70 else ""}</a>
                </div>
                """,
                unsafe_allow_html=True,
            )
    st.stop()

# ─── Validate ────────────────────────────────────────────────────────────────
if not search_term.strip():
    st.error("⚠️ Por favor, insira um termo de busca no painel lateral.")
    st.stop()

if not selected_countries:
    st.error("⚠️ Selecione ao menos um país.")
    st.stop()

# ─── Build results ───────────────────────────────────────────────────────────
term = search_term.strip()
results = {}

st.markdown(
    f'<p class="section-title">Resultados para "{term}"</p>',
    unsafe_allow_html=True,
)

# Chips
chips = ""
for c in selected_countries:
    chips += f'<span class="chip">{COUNTRIES[c]["flag"]} {c}</span>'
st.markdown(chips, unsafe_allow_html=True)
st.markdown("<br>", unsafe_allow_html=True)

cols = st.columns(2)
for i, country in enumerate(selected_countries):
    url = build_url(country, term)
    results[country] = url
    with cols[i % 2]:
        st.markdown(
            f"""
            <div class="country-card">
              <div class="country-name">{COUNTRIES[country]['flag']} {country}</div>
              <a href="{url}" target="_blank" class="country-url">🔗 Abrir pesquisa no parlamento</a><br><br>
              <code style="font-size:0.75rem;color:#888;word-break:break-all;">{url[:100]}{"…" if len(url) > 100 else ""}</code>
            </div>
            """,
            unsafe_allow_html=True,
        )

# ─── Metrics ─────────────────────────────────────────────────────────────────
st.markdown("<br>", unsafe_allow_html=True)
m1, m2 = st.columns(2)
m1.metric("🌍 Países pesquisados", len(selected_countries))
m2.metric("🔑 Termo de busca", f'"{term}"')

# ─── Word export ──────────────────────────────────────────────────────────────
def generate_word(term, results):
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Helper: set paragraph shading ---
    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # --- Header bar (table trick for background) ---
    header_table = doc.add_table(rows=1, cols=1)
    header_table.style = "Table Grid"
    cell = header_table.cell(0, 0)
    set_cell_bg(cell, "FF6B00")
    cell.width = Inches(6.5)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("⚖  LexGlobal")
    run.font.name = "Georgia"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p2 = cell.add_paragraph("Pesquisa de Legislação Internacional")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.runs[0]
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xFF, 0xE0, 0xC0)

    doc.add_paragraph()

    # --- Metadata block ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def add_meta(label, value):
        r = meta.add_run(f"{label}: ")
        r.font.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0xFF, 0x6B, 0x00)
        r2 = meta.add_run(value + "     ")
        r2.font.size = Pt(10.5)

    add_meta("Termo de Busca", f'"{term}"')
    add_meta("Países", str(len(results)))
    add_meta("Gerado em", datetime.now().strftime("%d/%m/%Y %H:%M"))

    # --- Orange divider ---
    div = doc.add_paragraph()
    div_run = div.add_run("─" * 80)
    div_run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x00)
    div_run.font.size = Pt(8)

    doc.add_paragraph()

    # --- Results table ---
    heading = doc.add_paragraph()
    h_run = heading.add_run("RESULTADOS DA PESQUISA")
    h_run.font.name = "Georgia"
    h_run.font.size = Pt(13)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    heading.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for idx, txt in enumerate(["PAÍS", "LINK DE BUSCA"]):
        set_cell_bg(hdr_cells[idx], "FF6B00")
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(txt)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, (country, url) in enumerate(results.items()):
        row = table.add_row()
        # alternating row bg
        bg = "FFF3E0" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            set_cell_bg(cell, bg)

        # Country cell
        cp = row.cells[0].paragraphs[0]
        flag = COUNTRIES[country]["flag"]
        r_flag = cp.add_run(f"{flag} ")
        r_flag.font.size = Pt(12)
        r_name = cp.add_run(country)
        r_name.font.bold = True
        r_name.font.size = Pt(10.5)
        r_name.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # URL cell – add hyperlink
        url_cell = row.cells[1]
        url_p = url_cell.paragraphs[0]
        # Add hyperlink via XML
        r_id = url_p.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "FF6B00")
        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), "18")  # 9pt
        rPr.append(color_el)
        rPr.append(u_el)
        rPr.append(sz_el)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = url
        t.set(qn("xml:space"), "preserve")
        new_run.append(t)
        hyperlink.append(new_run)
        url_p._p.append(hyperlink)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()

    # --- Footer ---
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(f"LexGlobal · Pesquisa de Legislação Internacional · {datetime.now().strftime('%d/%m/%Y')}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ─── Export button ────────────────────────────────────────────────────────────
st.markdown("<br>", unsafe_allow_html=True)
st.markdown('<p class="section-title">Exportar Resultados</p>', unsafe_allow_html=True)

word_bytes = generate_word(term, results)

filename = f"LexGlobal_{term.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

st.download_button(
    label="⬇️ Baixar relatório em Word (.docx)",
    data=word_bytes,
    file_name=filename,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
